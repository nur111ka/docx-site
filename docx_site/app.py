import os
import json
import re
from flask import Flask, render_template, request, send_from_directory, redirect, url_for, session, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from functools import wraps
from datetime import datetime

from config import Config
from database import (
    init_db, 
    # Функции для шаблонов
    create_template, delete_template, get_all_templates,
    get_template_fields, add_field_to_template, delete_field_from_template,
    update_field_in_template, save_template_replacements, get_template_replacements,
    # Функции для ключей
    generate_key, check_key, increment_usage, get_all_keys,
    deactivate_key, get_key_info, get_usage_stats, check_rate_limit
)

# ===== СОЗДАНИЕ ПРИЛОЖЕНИЯ =====
app = Flask(__name__)
app.secret_key = Config.SECRET_KEY

# Создаем папки
os.makedirs(Config.TEMPLATES_STORAGE, exist_ok=True)
os.makedirs(Config.OUTPUT_FOLDER, exist_ok=True)

# Инициализируем БД
init_db()

# ===== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ =====

def get_client_ip():
    """Получает IP адрес клиента."""
    if request.environ.get('HTTP_X_FORWARDED_FOR'):
        return request.environ.get('HTTP_X_FORWARDED_FOR').split(',')[0]
    return request.remote_addr

def replace_text_in_doc(doc, replacements_dict):
    """Заменяет текст в документе во всех местах (исправленная версия)."""
    
    def replace_in_text(text):
        """Заменяет все вхождения в тексте."""
        if not text:
            return text
        
        result = text
        for old, new in replacements_dict.items():
            if old in result:
                # Используем re.sub для гарантированной замены всех вхождений
                result = re.sub(re.escape(old), lambda m: new, result)
        return result
    
    def process_paragraph(paragraph):
        """Обрабатывает один параграф."""
        original_text = paragraph.text
        new_text = replace_in_text(original_text)
        
        if new_text != original_text:
            # Очищаем все runs
            for run in paragraph.runs:
                run.text = ""
            
            # Добавляем новый текст в первый run
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                # Сохраняем шрифт Times New Roman
                paragraph.runs[0].font.name = 'Times New Roman'
                paragraph.runs[0].font.size = None  # Сохраняем размер из шаблона
    
    # 1. Обрабатываем все обычные параграфы
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    
    # 2. Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    
    # 3. Обрабатываем верхние колонтитулы
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                process_paragraph(paragraph)
        
        if section.first_page_header:
            for paragraph in section.first_page_header.paragraphs:
                process_paragraph(paragraph)
    
    # 4. Обрабатываем нижние колонтитулы
    for section in doc.sections:
        if section.footer:
            for paragraph in section.footer.paragraphs:
                process_paragraph(paragraph)
        
        if section.first_page_footer:
            for paragraph in section.first_page_footer.paragraphs:
                process_paragraph(paragraph)
    
    # 5. Обрабатываем текстовые боксы (если есть)
    for shape in doc.inline_shapes:
        if hasattr(shape, 'text_frame'):
            for paragraph in shape.text_frame.paragraphs:
                process_paragraph(paragraph)

# ===== АДМИН ДЕКОРАТОР =====

def admin_required(f):
    """Декоратор для проверки авторизации админа."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'admin_auth' not in session:
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated_function

# ===== АДМИН МАРШРУТЫ =====

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    """Страница входа в админку."""
    if request.method == 'POST':
        password = request.form.get('password', '')
        if password == Config.ADMIN_PASSWORD:
            session['admin_auth'] = True
            return redirect(url_for('admin'))
        else:
            return render_template('admin_login.html', error="❌ Неверный пароль")
    
    return render_template('admin_login.html')

@app.route('/admin', methods=['GET', 'POST'])
@admin_required
def admin():
    """Админ-панель."""
    message = None
    error = None
    tab = request.args.get('tab', 'templates')
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'create_key':
            template_name = request.form.get('key_client_name', '').strip()
            limit_count = request.form.get('limit_count', '10')
            
            if not template_name:
                error = "❌ Выберите шаблон"
            else:
                try:
                    limit_count = int(limit_count)
                    if limit_count < 1:
                        error = "❌ Лимит должен быть больше 0"
                    else:
                        new_key = generate_key(template_name, limit_count)
                        message = f"✅ Ключ создан: <code style='background: #f0f0f0; padding: 5px;'>{new_key}</code>"
                        tab = 'keys'
                except ValueError:
                    error = "❌ Лимит должен быть числом"
        
        elif action == 'deactivate_key':
            api_key = request.form.get('api_key', '')
            if api_key:
                deactivate_key(api_key)
                message = "✅ Ключ деактивирован"
                tab = 'keys'
    
    # Получаем данные
    keys = get_all_keys()
    templates = get_all_templates()
    stats = get_usage_stats()
    
    return render_template('admin.html',
                         tab=tab,
                         keys=keys,
                         templates=templates,
                         stats=stats,
                         message=message,
                         error=error,
                         get_all_templates=get_all_templates)

@app.route('/admin/templates/manage', methods=['GET', 'POST'])
@admin_required
def manage_templates():
    """Управление шаблонами."""
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'create_template':
            template_name = request.form.get('template_name', '').strip()
            display_name = request.form.get('display_name', '').strip()
            
            if template_name and display_name:
                if create_template(template_name, display_name):
                    # Создаем дефолтные поля
                    add_field_to_template(template_name, 'client_name', 'ФИО клиента', 'text')
                    add_field_to_template(template_name, 'client_phone', 'Телефон', 'text')
                    
                    # Создаем дефолтный JSON
                    default_json = {
                        "[COMPANY_NAME]": "Ваша компания",
                        "[COMPANY_ADDRESS]": "г. Москва",
                        "[TODAY_DATE]": datetime.now().strftime('%Y-%m-%d')
                    }
                    save_template_replacements(template_name, json.dumps(default_json, ensure_ascii=False))
                    
                    return redirect(url_for('manage_template', template_name=template_name))
        
        elif action == 'delete_template':
            template_name = request.form.get('template_name')
            if template_name:
                delete_template(template_name)
    
    templates = get_all_templates()
    return render_template('manage_templates.html', templates=templates)

@app.route('/admin/template/<template_name>', methods=['GET', 'POST'])
@admin_required
def manage_template(template_name):
    """Управление конкретным шаблоном."""
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add_field':
            field_name = request.form.get('field_name', '').strip()
            field_label = request.form.get('field_label', '').strip()
            field_type = request.form.get('field_type', 'text')
            
            if field_name and field_label:
                add_field_to_template(template_name, field_name, field_label, field_type)
        
        elif action == 'delete_field':
            field_name = request.form.get('field_name')
            if field_name:
                delete_field_from_template(template_name, field_name)
        
        elif action == 'update_field':
            field_name = request.form.get('field_name')
            field_label = request.form.get('field_label', '').strip()
            field_type = request.form.get('field_type', 'text')
            
            if field_label:
                update_field_in_template(template_name, field_name, field_label, field_type)
        
        elif action == 'save_replacements':
            replacements_text = request.form.get('replacements_json', '{}')
            try:
                # Валидируем JSON
                json.loads(replacements_text)
                save_template_replacements(template_name, replacements_text)
                return jsonify({'success': True, 'message': 'JSON сохранен'})
            except Exception as e:
                return jsonify({'success': False, 'message': f'Ошибка JSON: {str(e)}'})
        
        elif action == 'upload_template_file':
            if 'template_file' in request.files:
                f = request.files['template_file']
                if f.filename.endswith('.docx'):
                    safe_name = secure_filename(template_name)
                    template_path = os.path.join(Config.TEMPLATES_STORAGE, f"{safe_name}.docx")
                    f.save(template_path)
                    return redirect(url_for('manage_template', template_name=template_name))
    
    # Получаем данные шаблона
    fields = get_template_fields(template_name)
    replacements_json = get_template_replacements(template_name)
    
    # Проверяем наличие файла шаблона
    template_file_exists = os.path.exists(
        os.path.join(Config.TEMPLATES_STORAGE, f"{secure_filename(template_name)}.docx")
    )
    
    return render_template('manage_template.html',
                         template_name=template_name,
                         fields=fields,
                         replacements_json=replacements_json,
                         template_file_exists=template_file_exists)

# ===== КЛИЕНТСКИЕ МАРШРУТЫ =====

@app.route('/', methods=['GET', 'POST'])
def client():
    """Клиентская форма для генерации документов."""
    api_key = request.args.get('key', '')
    fields = []
    error = None
    remaining = 0
    template_name = None
    client_name = None
    total_used = 0
    percentage = 0
    limit_count = 0
    
    # GET запрос
    if request.method == 'GET' and api_key:
        valid, msg = check_key(api_key)
        if valid:
            template_name = msg
            fields = get_template_fields(template_name)
            key_info = get_key_info(api_key)
            if key_info:
                limit_count, used_count, _, status = key_info
                total_used = used_count
                remaining = limit_count - used_count
                if limit_count > 0:
                    percentage = min((used_count / limit_count * 100), 100)
        else:
            error = msg
            api_key = ""
    
    # POST запрос
    elif request.method == 'POST':
        api_key = request.form.get('api_key', '').strip()
        client_ip = get_client_ip()
        
        # Проверяем rate limit
        rate_ok, rate_msg = check_rate_limit(api_key, client_ip, 
                                            Config.RATE_LIMIT_REQUESTS, 
                                            Config.RATE_LIMIT_PERIOD)
        if not rate_ok:
            increment_usage(api_key, client_ip, "rate_limit_exceeded", rate_msg)
            return render_template('client.html',
                                 error=rate_msg,
                                 api_key=api_key)
        
        valid, msg = check_key(api_key)
        if not valid:
            return render_template('client.html', error=msg, api_key=api_key)
        
        template_name = msg
        key_info = get_key_info(api_key)
        if not key_info:
            return render_template('client.html', error="❌ Ключ не найден", api_key=api_key)
        
        limit_count, used_count, _, status = key_info
        remaining = limit_count - used_count
        total_used = used_count
        if limit_count > 0:
            percentage = min((used_count / limit_count * 100), 100)
        
        # Проверяем шаблон
        safe_name = secure_filename(template_name)
        template_file = os.path.join(Config.TEMPLATES_STORAGE, f"{safe_name}.docx")
        
        if not os.path.exists(template_file):
            fields = get_template_fields(template_name)
            return render_template('client.html',
                                 error=f"❌ Файл шаблона не найден",
                                 api_key=api_key,
                                 fields=fields,
                                 remaining=remaining,
                                 template_name=template_name,
                                 client_name=client_name,
                                 total_used=total_used,
                                 percentage=percentage,
                                 limit_count=limit_count)
        
        # Собираем замены
        replacements = {}
        
        # 1. Замены из JSON базы данных
        try:
            replacements_json = get_template_replacements(template_name)
            replacements.update(json.loads(replacements_json))
        except Exception as e:
            print(f"Ошибка загрузки JSON замен: {e}")
        
        # 2. Замены из формы пользователя
        fields_list = get_template_fields(template_name)
        for field_name, field_label, field_type, field_order in fields_list:
            value = request.form.get(field_name, '')
            if value:
                # Добавляем скобки к имени поля для замены
                replacements[f"[{field_name}]"] = value
        
        # 3. Добавляем системные замены
        replacements["[TODAY_DATE]"] = datetime.now().strftime('%d.%m.%Y')
        replacements["[TODAY_DATE_FULL]"] = datetime.now().strftime('%d %B %Y года')
        
        # Генерируем документ
        try:
            print(f"Начинаем обработку документа с {len(replacements)} заменами")
            print(f"Замены: {replacements}")
            
            doc = Document(template_file)
            
            # Подсчитываем количество параграфов для отладки
            total_paragraphs = len(doc.paragraphs)
            print(f"Документ загружен. Параграфов: {total_paragraphs}")
            
            # Применяем замены
            replace_text_in_doc(doc, replacements)
            
            # Генерируем имя файла
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            # Получаем имя клиента для имени файла
            client_name_value = request.form.get('client_name', '')
            if client_name_value:
                safe_client_name = secure_filename(client_name_value[:50])
                output_filename = f"{template_name}_{safe_client_name}_{timestamp}.docx"
            else:
                output_filename = f"{template_name}_{timestamp}.docx"
            
            output_path = os.path.join(Config.OUTPUT_FOLDER, output_filename)
            
            # Сохраняем документ
            doc.save(output_path)
            print(f"Документ сохранен: {output_path}")
            
            # Увеличиваем счетчик использования
            increment_usage(api_key, client_ip, "success", output_filename)
            
            # Отправляем файл клиенту
            return send_from_directory(Config.OUTPUT_FOLDER,
                                     output_filename,
                                     as_attachment=True,
                                     download_name=output_filename)
        
        except Exception as e:
            print(f"❌ Ошибка при обработке документа: {str(e)}")
            import traceback
            traceback.print_exc()
            
            increment_usage(api_key, client_ip, "error", str(e))
            fields = get_template_fields(template_name)
            return render_template('client.html',
                                 error=f"❌ Ошибка при обработке документа: {str(e)}",
                                 api_key=api_key,
                                 fields=fields,
                                 remaining=remaining,
                                 template_name=template_name,
                                 client_name=client_name,
                                 total_used=total_used,
                                 percentage=percentage,
                                 limit_count=limit_count)
    
    return render_template('client.html',
                         api_key=api_key,
                         fields=fields,
                         error=error,
                         remaining=remaining,
                         template_name=template_name,
                         client_name=client_name,
                         total_used=total_used,
                         percentage=percentage,
                         limit_count=limit_count)

@app.route('/admin/logout')
@admin_required
def admin_logout():
    """Выход из админки."""
    session.pop('admin_auth', None)
    return redirect(url_for('admin_login'))

# ===== ЗАПУСК СЕРВЕРА =====

if __name__ == '__main__':
    # Очистка экрана
    import os
    os.system('cls' if os.name == 'nt' else 'clear')
    
    print("=" * 60)
    print("🚀 DOCX Generator PRO - Запуск системы")
    print("=" * 60)
    print(f"📁 Папка шаблонов: {Config.TEMPLATES_STORAGE}")
    print(f"📁 Выходные файлы: {Config.OUTPUT_FOLDER}")
    print(f"🗄️ База данных: {Config.DATABASE_PATH}")
    print("=" * 60)
    print("🔐 Админ доступ:")
    print(f"  • Пароль: {Config.ADMIN_PASSWORD}")
    print(f"  • URL: http://localhost:5000/admin/login")
    print("=" * 60)
    print("🌐 Клиентский доступ:")
    print("  • Главная: http://localhost:5000")
    print("  • Форма: http://localhost:5000/?key=ВАШ_КЛЮЧ")
    print("=" * 60)
    print("⚡ Система готова к работе!")
    print("=" * 60)
    print("\n📋 Логи сервера:")
    print("-" * 60)
    
    try:
        # Проверяем доступность порта
        import socket
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(1)
        result = sock.connect_ex(('localhost', 5000))
        sock.close()
        
        if result == 0:
            print("⚠️  Внимание: Порт 5000 уже занят!")
            print("   Возможно, сервер уже запущен в другом окне.")
            print("   Закройте предыдущую сессию или используйте другой порт.")
            print("\n   Для смены порта измените в app.py:")
            print("   app.run(debug=True, host='0.0.0.0', port=5000)")
            print("   на")
            print("   app.run(debug=True, host='0.0.0.0', port=5001)")
            print("-" * 60)
    except:
        pass
    
    # Запуск Flask сервера
    app.run(debug=True, host='0.0.0.0', port=5000)